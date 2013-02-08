#!/usr/bin/python

# word2lyx is a document parsing script used to 
# convert Microsoft Word documents to LyX documents.
# (C) Robert Oakes, 2012. Released under the terms
# of the GNU Lesser General Public License (LGPL).
# Preferred Dependencies: lxml, elyxer

# Import docx library and supporting tools
import math
import sys 
import os
import tempfile
import shutil
import codecs
import urllib
import urlparse

# Import docx parsing classes
from docx import read as docxread
from docx import properties as docx

# Import lyx parsing classes
from lyx.tables import *
from lyx.charstyles import *
#  import lyx.templates as templates

# Import NTI things
from nti.contentfragments import interfaces as frg_interfaces
from nti.contentfragments.latex import PlainTextToLatexFragmentConverter

##------------ Setup Input File, Output File, and Template ------------## 
#SCRIPT_DIR = os.path.abspath(os.path.dirname(__file__))
#template = os.path.join(SCRIPT_DIR, 'lyx/templates/article.w2l')

# Get input and output file
try:
	inputfile = sys.argv[1]
	# Check to see if the file exists
	if not os.path.exists(inputfile):
		print 'Error: Unable to locate input file'
		exit()
	# Check to see if the file is a docx file
	basename = os.path.basename(inputfile)
	fileext = basename.split('.')[len(basename.split('.'))-1]
	if fileext != 'docx':
		print 'Error: Invalid input file. word2lyx only supports docx files.'
		exit()
	outputfile = sys.argv[2]

# If there is an error encoutered, return general error message
except:
	print 'lyx2word encountered an error with the file' 
	exit()

##------------ Global Structures ------------## 

CHARSTYLE_LIST = []
PROC_REL = []
DOC_REL = docxread.openDocxRelationships(inputfile)
END_REL = docxread.openDocxEndnotesRelationships(inputfile)

# Load footnotes
DOC_FOOTNOTES = docxread.openDocxFootnotes(inputfile)

# Load endnotes
DOC_ENDNOTES = docxread.openDocxEndnotes(inputfile)

#PARAGRAPH_STYLES = doc_options.returnStyles('ParagraphStyles')
#TABLE_STYLES = doc_options.returnStyles('TableStyles')
#CHARACTER_STYLES = doc_options.returnStyles('CharacterStyles')
#IGNORE_STYLES = doc_options.returnSetting('IgnoreStyles')

#IMAGE_FOLDER = doc_options.returnSetting('ImageDir')
IMAGE_FOLDER = 'Images'


##------------ Class Methods ------------## 

class _Container(frg_interfaces.LatexContentFragment):

        children = ()

        def add_child( self, child ):
                if self.children == ():
                        self.children = []
                self.children.append( child )

class _WrappedElement(_Container):
        wrapper = None

        def __new__( cls, text ):
                return super(_WrappedElement,cls).__new__( cls, '\\' + cls.wrapper + '{' + text + '}' )

        def __init__( self, text=None ):
                # Note: __new__ does all the actual work, because these are immutable as strings
                super(_WrappedElement,self).__init__( self, '\\' + self.wrapper + '{' + text + '}' )

class _TextNode(frg_interfaces.LatexContentFragment):
	pass

        def __new__( cls, text ='' ):
                return super(_TextNode,cls).__new__( cls, PlainTextToLatexFragmentConverter(text) )

        def __init__( self, text='' ):
                # Note: __new__ does all the actual work, because these are immutable as strings
                super(_TextNode,self).__init__( self, PlainTextToLatexFragmentConverter(text) )


class _Footnote(_WrappedElement):
        wrapper = 'footnote'

class _Chapter(_WrappedElement):
        wrapper = 'chapter'

class _Section(_WrappedElement):
        wrapper = 'section'

class _SubSection(_WrappedElement):
        wrapper = 'subsection'

class _SubSubSection(_WrappedElement):
        wrapper = 'subsubsection'

class _Paragraph(_WrappedElement):
        wrapper = 'paragraph'

class _SubParagraph(_WrappedElement):
        wrapper = 'subparagraph'

class _SubSubParagraph(_WrappedElement):
        wrapper = 'subsubparagraph'

class _Label(_WrappedElement):
        wrapper = 'label'

class _Title(_WrappedElement):
        wrapper = 'title'

class _Author(_WrappedElement):
        wrapper = 'author'

class _TextIT(_WrappedElement):
        wrapper = 'textit'

class _TextBF(_WrappedElement):
        wrapper = 'textbf'

class _Uline(_WrappedElement):
        wrapper = 'uline'

class _NTIIncludeVideo(_WrappedElement):
        wrapper = 'ntiincludevideo'

class _NTIImageHref(_WrappedElement):
        wrapper = 'ntiimagehref'

class _Newline(_Container):

        def __new__( cls ):
                return super(_Newline,cls).__new__( cls, '\\newline\n' )

        def __init__( self ):
                super(_Newline,self).__init__( self, '\\newline\n' )

class _href(_Container):

        def __new__( cls, url, text=None ):
		if text:
			_t = url + '}{' + text
		else:
			_t = url
                return super(_href,cls).__new__( cls, '\\href{' + _t + '}' )

        def __init__( self, url, text=None ):
                # Note: __new__ does all the actual work, because these are immutable as strings
		if text:
			_t = url + '}{' + text
		else:
			_t = url
                super(_href,self).__init__( self, '\\href{' + _t + '}' )

class _Image(_Container):

        def __new__( cls, image_file, parms=None ):
		if parms:
			_t = '[%s]{%s}' % (parms,image_file)
		else:
			_t = '{%s}' % image_file
                return super(_Image,cls).__new__( cls, '\\includegraphics' + _t  )

        def __init__( self, image_file, parms=None ):
                # Note: __new__ does all the actual work, because these are immutable as strings
		if parms:
			_t = '[%s]{%s}' % (parms,image_file)
		else:
			_t = '{%s}' % image_file
                super(_Image,self).__init__( self, '\\includegraphics' + _t  )

# These classes encapsulate WordprocessingML structures
class _DocxStructureElement( object ):

	def __init__( self ):
		self.styles = []
		self.children = []

	def __str__( self ):
		val = u''
		for child in self.children:
			if isinstance( child, _DocxStructureElement ):
				val = val + str(child)
			else:
				val = val + child

		return val

	def raw( self ):
		val = u''
		for child in self.children:
			if isinstance( child, _DocxStructureElement ):
				val = val + child.raw()
			else:
				val = val + child
		return val

	def addChild( self, elem ):
		self.children.append(elem)

	def removeChild( self, elem ):
		self.children.remove( elem )

	def addStyle( self, style ):
		self.styles.append(style)

	def removeStyle( self, style ):
		self.styles.remove(style)


class _TextRun( _DocxStructureElement ):
	STYLES = { 'bold': _TextBF,
		   'italic': _TextIT,
		   'underline': _Uline}

	def __str__( self ):
		val = super(_TextRun, self).__str__()

		for style in self.styles:
			if val:
				if style in self.__class__.STYLES:
					val = self.__class__.STYLES[style](val)
				else:
					print('Unhandled style: %s' % style)

		return val

class _Paragraph( _TextRun ):
	STYLES = { 'Heading1': _Chapter,
		   'Heading2': _Section,
		   'Heading3': _SubSection,
		   'Heading4': _SubSubSection,
		   'Heading5': _Paragraph,
		   'Heading6': _SubParagraph,
		   'Heading7': _SubSubParagraph}

	def __str__( self ):
		return '\n' + super(_Paragraph, self).__str__() + '\n'

CONTAINERS = { 'blockquote': 'quote',
	       'center': 'center' }

IGNORED_TAGS = [ '{'+docx.nsprefixes['w']+'}ind',
		 '{'+docx.nsprefixes['w']+'}sectPr',
		 '{'+docx.nsprefixes['w']+'}proofErr',
		 '{'+docx.nsprefixes['w']+'}noProof',
		 '{'+docx.nsprefixes['w']+'}commentReference',
		 '{'+docx.nsprefixes['w']+'}commentRangeEnd',
		 '{'+docx.nsprefixes['w']+'}commentRangeStart',
		 '{'+docx.nsprefixes['w']+'}bookmarkEnd',
		 '{'+docx.nsprefixes['w']+'}bookmarkStart',
		 '{'+docx.nsprefixes['w']+'}lastRenderedPageBreak']

PACKAGES = [ 'graphicx',
	     'ntilatexmacros',
	     'hyperref',
	     'ulem' ]

def processDocument(inputfile, outputfile):
	'''Method that parses the input document and translates it to a valid LyX structure. 
		Returns a text string with XML structures translated to equivalent LyX insets. 
		It uses the layouts in article and book as a base. If there are character styles 
		not defined in article or book, it will create placeholder styles using flex:inset.
		These may be modified at a later point by the user.'''

	# Open document from input string
	print "Beginning Conversion of " + inputfile
	document = docxread.opendocx(inputfile)

	# Create output strings
	doc_header = u''
	doc_layout = u''
	doc_body = u''
	
	# Iterate over the structure of the document, process document body
	for element in document.iterchildren():
		
		# Process Elements in Document Body
		if element.tag == '{'+docx.nsprefixes['w']+'}body':
			doc_body = doc_body + u'\\begin{document}\n'
			doc_body = doc_body + processBody( element )
			doc_body = doc_body + u'\\end{document}\n'

	# Copy Document Images to Subfolder
	if len(PROC_REL) > 0:
		img_exportfolder = os.path.join(os.path.abspath(os.path.dirname(outputfile)),
			IMAGE_FOLDER)
		docxread.getDocumentImages(inputfile, PROC_REL, img_exportfolder)
	
	# Retrieve character style information for document
	doc_styles = docxread.openDocxStyles(inputfile)
	# Retrieve font information for document
	doc_fonts = docxread.openFontTable(inputfile)

	# Create the document header
	# Set the document class
	doc_header = u'\\documentclass{book}\n'

	# Add packages
	for package in PACKAGES:
		doc_header = doc_header + '\\usepackage{' + package + '}\n'

	doc_header = doc_header

	# Combine the document header and body
	doc = doc_header + doc_body

	return doc

def processBody( body, rels = DOC_REL ):
	"""Process the content of a WordprocessingML body tag."""
	body_text = u''

	for element in body.iterchildren():

		# P (paragraph) Elements
		if element.tag == '{'+docx.nsprefixes['w']+'}p':
			body_text = body_text + str(processParagraph(element))

		elif element.tag == '{'+docx.nsprefixes['w']+'}tbl':
			body_text = body_text + processTable(element)

		# Skip elements in IGNORED_TAGS
		elif element.tag in IGNORED_TAGS:
			pass

		else:
			print('Did not handle body element: %s' % element.tag)

	return body_text

def processParagraph(paragraph, lyxTable = False, rels = DOC_REL ):
	'''Processes the text of a given paragraph into insets and text.'''
	
	me = _Paragraph()
	fields = []
	# Scan the elements in the paragraph and extract information
	for element in paragraph.iterchildren():

		# Process Text Runs
		if element.tag == '{'+docx.nsprefixes['w']+'}r':
			me.addChild(processTextRun(element, fields = fields, rels = rels))
			image_text = processRelationship(element)
			if len(image_text) > 0:
				me.addChild(image_text)
		
		# Look for hyperlinks
		elif (element.tag == '{'+docx.nsprefixes['w']+'}hyperlink'):
			me.addChild(processHyperlink(element, rels = rels))

		# Paragraph Properties
		elif element.tag == '{'+docx.nsprefixes['w']+'}pPr':
			for sub_element in element.iterchildren():
				# Look for Paragraph Styles
				if sub_element.tag == '{'+docx.nsprefixes['w']+'}pStyle':
					me.addStyle(sub_element.attrib['{'+docx.nsprefixes['w']+'}val'])

				# We don't care about the paragraph mark character in LaTeX so ignore formattingi it.
				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}rPr':
					pass

				# Skip elements in IGNORED_TAGS
				elif sub_element.tag in IGNORED_TAGS:
					pass

				else:
					print('Unhandled paragraph property: %s' % sub_element.tag)

		# Skip elements in IGNORED_TAGS
		elif element.tag in IGNORED_TAGS:
			pass

		# We did not handle the element
		else:
			print('Did not handle paragraph element: %s' % element.tag)
	
	return me


def processTextRun(textrun, fields = [], rels = DOC_REL ):
	'''Process a paragraph textrun, parse for character styles'''

	me = _TextRun()

	for element in textrun.iterchildren():

		# Look for run properties
		if element.tag == '{'+docx.nsprefixes['w']+'}rPr':
			for sub_element in element.iterchildren():
				if sub_element.tag == '{'+docx.nsprefixes['w']+'}rStyle':
					me.addStyle(sub_element.attrib['{'+docx.nsprefixes['w']+'}val'])

				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}b':
                                        me.addStyle('bold')

				# Ignore bold complex script specification
				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}bCs':
                                        pass

				# Ignore run content color specification
				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}color':
                                        pass

				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}i':
                                        me.addStyle('italic')

				# Ignore italic  complex script specification
				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}iCs':
                                        pass

				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}u':
                                        me.addStyle('underline')

				# Ignore run level font specifications
				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}rFonts':
                                        pass

				# Ignore run level font size specifications
				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}sz':
                                        pass

				# Ignore run level complex script font size specifications
				elif sub_element.tag == '{'+docx.nsprefixes['w']+'}szCs':
                                        pass

				# Skip elements in IGNORED_TAGS
				elif sub_element.tag in IGNORED_TAGS:
					pass

				else:
					print('Unhandled run property: %s' % sub_element.tag)

		# Find run text
		elif (element.tag == '{'+docx.nsprefixes['w']+'}t'): 
			# If not character style, append to the end of the paragraph
			if element.text:
				me.addChild( _TextNode(element.text) )

		# Look for hyperlinks
		elif (element.tag == '{'+docx.nsprefixes['w']+'}hyperlink'):
			me.addChild( processHyperlink(element, rels = rels) )

		# Look for complex fields
		elif (element.tag == '{'+docx.nsprefixes['w']+'}fldChar'):
			type = element.attrib['{'+docx.nsprefixes['w']+'}fldCharType']
			if 'begin' in type:
				# Push the element onto the field stack
				fields.append(element)
			elif 'end' in type:
				elems = []
				# Pop the element off of the field stack
				while fields:
					if not isinstance( fields[len(fields)-1], _TextRun ) \
						    and  '{'+docx.nsprefixes['w']+'}fldChar' in fields[len(fields)-1].tag:
						if 'begin' in fields[len(fields)-1].attrib['{'+docx.nsprefixes['w']+'}fldCharType']:
							# We have found the begining of the field so pop it off and
							# stop searching
							fields.pop()
							break;
						else:
							fields.pop()
					else:
						elems.append(fields.pop())
				me.addChild( processComplexField( elems, rels = rels ) )
			elif 'separate' in type:
				# This node was used to separate the field "command" from the result. It does not seem
				# to make sense with how we are processing the document so it is being ignored.
				pass

		# Look for field codes
		elif (element.tag == '{'+docx.nsprefixes['w']+'}instrText'):
			fields.append(element)

		# Look for footnotes
		elif (element.tag == '{'+docx.nsprefixes['w']+'}footnoteReference'):
			me.addChild( processFootnote(element) )

		# Look for endnotes
		elif (element.tag == '{'+docx.nsprefixes['w']+'}endnoteReference'):
			me.addChild( processEndnote(element) )

		# Look for endnotes
		elif (element.tag == '{'+docx.nsprefixes['w']+'}br'):
			me.addChild( _Newline() )

		# Look for other runs embedded in this run, process recursively
		elif (element.tag == '{'+docx.nsprefixes['w']+'}r'):
			me.addChild( processTextRun(element, fields = fields, rels = rels) )

		# Skip elements in IGNORED_TAGS
		elif element.tag in IGNORED_TAGS:
			pass

		# We did not handle the element
		else:
			print('Did not handle run element: %s' % element.tag)

	if fields:
		fields.append( me )
		return ''
	else:
		return me

def _buildHref( url, text ):
	if 'YouTube Video' in text:
		scheme = ''
		netloc = 'www.youtube.com'
		path = '/embed'
		query = { 'html5': '1', 'rel': '0' }
		parsed_url = urlparse.urlsplit( url )

		if len(parsed_url.path.split('/')) > 2:
			print(parsed_url.path)
			#Then assume that we were passed a more complete URL
			path = parsed_url.path
		else:
			#Assume we were given a 'shortened' URL.
			path = path + parsed_url.path

		# Add any query args to ours
		query.update( urlparse.parse_qsl(parsed_url.query) )
		target = urlparse.urlunsplit( (scheme, netloc, path, urllib.urlencode(query), '') )
		text = _NTIIncludeVideo( target )
	elif 'Thumbnail' in text:
		text = _NTIImageHref( url )
	else:
		text = _href( url, text)

	return text

def processHyperlink(node, rels = DOC_REL ):
	'''Process a hyperlink element'''

	text = ''

	rId = node.attrib['{'+docx.nsprefixes['r']+'}id']
	rel_type, rel_target = relationshipProperties(rId, rels)

	for element in node.iterchildren():
		# Look for footnotes
		if (element.tag == '{'+docx.nsprefixes['w']+'}footnoteReference'):
			text = text + processFootnote(element)

		# Look for endnotes
		if (element.tag == '{'+docx.nsprefixes['w']+'}endnoteReference'):
			text = text + processEndnote(element)

		# Look for embedded runs
		if (element.tag == '{'+docx.nsprefixes['w']+'}r'):
			_t = processTextRun(element, rels = rels).raw()
			text = text + _t

	if text:
		text = _buildHref( rel_target, text )
	else:
		text = ''

	return text

def processField( field, result, rels=DOC_REL ):
	text = _TextRun()
	field_code = field.split()

	if field_code[0] == 'HYPERLINK':
		text.addChild( _buildHref( field_code[len(field_code)-1].replace('"',''), str(result) ) )
	else:
		print( (field, str(result) ) )
	return text

def processComplexField( elements, rels=DOC_REL ):
	field = ''
	result = _TextRun()
	for element in elements:
		if isinstance( element, _TextRun ):
			result.addChild( element )
		elif element.tag == '{'+docx.nsprefixes['w']+'}instrText':
			field = field + _TextNode( element.text )
	return processField( field, result )

def processFootnote(footnote_ref):
	'''Locate footnote and write a footnote inset.'''

	# Retrieve the footnote Id Number
	id = footnote_ref.attrib['{'+docx.nsprefixes['w']+'}id']
	text = u''

	# Retrieve the footnote text
	for footnote in DOC_FOOTNOTES.iterchildren():
		if footnote.attrib['{'+docx.nsprefixes['w']+'}id'] == id:
			for foot_sub in footnote.iterchildren():
				# Process paragraphs found in the note
				if foot_sub.tag == '{'+docx.nsprefixes['w']+'}p':
					text = text + processParagraph(foot_sub, rels = FOOT_REL )

	# SAJ: In NTI LaTeX footnotes and endnotes are handled identically.
	return _Footnote( text )

def processEndnote(endnote_ref):
	'''Locate endnote and write a endnote tag.'''

	# Retrieve the endnote Id Number
	id = endnote_ref.attrib['{'+docx.nsprefixes['w']+'}id']
	text = u''

	# Retrieve the endnote text
	for endnote in DOC_ENDNOTES.iterchildren():
		if endnote.attrib['{'+docx.nsprefixes['w']+'}id'] == id:
			for end_sub in endnote.iterchildren():
				# Process paragraphs found in the note
				if end_sub.tag == '{'+docx.nsprefixes['w']+'}p':
					text = text + processParagraph(end_sub, rels = END_REL)

	# SAJ: In NTI LaTeX footnotes and endnotes are handled identically.
	return _Footnote( text )

def processRelationship(relationship):
	'''Process relationships in the document. '''
	rel_text = u''
	
	# Iterate through the relationship properties, process
	for element in relationship.iter():
		if element.tag == '{'+docx.nsprefixes['w']+'}drawing':
			rel_text = processImage(element)
	return rel_text


def processImage(image):
	'''Process images in the document. Specify that image should be copied from the 
		zip archive to a destination directory. Add a LyX image inset to the document. 
		The method ignores Microsoft specific charts and other files that have been 
		embedded.'''

	data = {}

	# Iterate through the image properties, process
	for element in image.iter():
		# Search for Image Properties, contained in blipFill
		if element.tag == '{'+docx.nsprefixes['pic']+'}blipFill':

			# Retrieve the Image rId and filename
			for imageprop in element.iter():
				if imageprop.tag == '{'+docx.nsprefixes['a']+'}blip':
					data['id'] = imageprop.attrib['{'+docx.nsprefixes['r']+'}embed']
					data['type'], data['target'] = relationshipProperties(data['id'])
					PROC_REL.append(os.path.basename(data['target']))
					break
		elif element.tag == '{'+docx.nsprefixes['wp']+'}inline':
			for imageprop in element.iter():
				if imageprop.tag == '{'+docx.nsprefixes['a']+'}ext':
					# Convert EMU to inches. 1 inch = 914400 EMU
					if 'cy' in imageprop.attrib:
						data['height'] = (float(imageprop.attrib['cy']) / 914400)
					if 'cx' in imageprop.attrib:
						data['width'] = (float(imageprop.attrib['cx']) / 914400)
			
	# SAJ: We are width constrained. If we only set the width attribute LaTeX will make sure 
	# we preserve aspect ratio.
	parms = None
	if 'width' in data:
		# SAJ: The constant 72 is for 72 dpi which is what ghostscript assumes for pixel / inch 
		# conversions. Because the number of pixels needs to be and integer, take the ceiling of
		# the result and then force to an int.
		width = int(math.ceil(data['width'] * 72))
		# SAJ: Cap width at our max width. Need to make this a constant variable.
		if width > 600:
			width = 600
		parms = 'width=%spx' % width

	# Write inset to file, ignore charts and other types of files
	if 'type' in data and data['type'] == docx.nsprefixes['wordimage']:
		text = '\n'  + \
		    _Image( os.path.join( IMAGE_FOLDER , os.path.splitext( os.path.basename(data['target']) )[0] ), 
			      parms) + \
		    '\n\n'
	else:
		text = ''

	return text


def relationshipProperties( rId, rels = DOC_REL ):
	'''Parse the document relationships to retrieve the type and target for a 
		specified document element.'''

	# Search through rels to retrieve relationship properties
	for relationship in rels.iter():
		try:
			if relationship.attrib['Id'] == rId:
				return relationship.attrib['Type'], relationship.attrib['Target']
		except:
			pass


def processTable(table_element):
	'''Parse a table element and translate to LyX.'''
	# Create a logical representation of the lyxTable
	lyx_table = lyxTable()
	
	# Parse element and determine table properties
	for element in table_element.iterchildren():

		# Parse table properties
		if element.tag == '{'+docx.nsprefixes['w']+'}tblPr':
			lyx_table = tableProperties(element, lyx_table)
				
		# Create column properties, add to lyxTable
		if element.tag == '{'+docx.nsprefixes['w']+'}tblGrid':
			lyx_table = tableGridProperties(element, lyx_table)
			
		# Create table rows, cells. Populate with data and properties.
		# Add populated rows to the lyxTable
		if element.tag == '{'+docx.nsprefixes['w']+'}tr':
			lyx_row = processTableRow(element)
			lyx_table.rows.append(lyx_row)
	
	# Add additional data to table structure
	lyx_table.numRows = len(lyx_table.rows)
	out_text = writeLyxTable(lyx_table)
	out_text = '\\begin_layout Standard\n' + out_text + '\\end_layout\n'
	
	return out_text


def tableProperties(property_list, lyx_table):
	'''Parse the general table properties, such as determining the location of table 
		borders, whether there is a grid in use, and the width of the table.'''

	# Scan the tblProperty elements
	for prop in property_list.iter():

		# Look for general table style settings
		if prop.tag == '{'+docx.nsprefixes['w']+'}tblStyle':
			if '{'+docx.nsprefixes['w']+'}val' in prop.attrib.keys():
				# Apply a grid to the table, activate all borders
				if prop.attrib['{'+docx.nsprefixes['w']+'}val'] == 'TableGrid':
					for border in lyx_table.borders.data.keys():
						lyx_table.borders.data[border] = 'true'

		# Look for table border settings
		if prop.tag == '{'+docx.nsprefixes['w']+'}tblBorders':

			# Apply settings to the border settings for the lyxTable object
			for border in prop.iterchildren():

				# Top Border
				if border.tag == '{'+docx.nsprefixes['w']+'}top':
					lyx_table.borders.data['top'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

				# Bottom Border
				if border.tag == '{'+docx.nsprefixes['w']+'}bottom':
					lyx_table.borders.data['bottom'] = border.attrib['{'+docx.nsprefixes['w']+'}val']	
				# Left Border
				if border.tag == '{'+docx.nsprefixes['w']+'}left':
					lyx_table.borders.data['left'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

				# Right Border
				if border.tag == '{'+docx.nsprefixes['w']+'}right':
					lyx_table.borders.data['right'] = border.attrib['{'+docx.nsprefixes['w']+'}val']
				
				# Inside Horizontal Borders
				if border.tag == '{'+docx.nsprefixes['w']+'}insideH':
					lyx_table.borders.data['insideH'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

				# Inside Veritcal Border
				if border.tag == '{'+docx.nsprefixes['w']+'}insideV':
					lyx_table.borders.data['insideV'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

		# Table Width and Units
		if prop.tag == '{'+docx.nsprefixes['w']+'}tblW':
			lyx_table.width = prop.attrib['{'+docx.nsprefixes['w']+'}w']
			lyx_table.units = prop.attrib['{'+docx.nsprefixes['w']+'}type']

	return lyx_table
	

def tableGridProperties(table_grid, lyx_table):
	'''Parse grid column properties. Looks at the table column properties to set width and text styles. Takes both the element to be parsed and a lyxTable object as input. Returns the lyxTable with column objects appended.'''
	
	# Scan the tblGrid Element for gridCol elements, retrieve relative width for each
	for table_col in table_grid.iter():
		if table_col.tag == '{'+docx.nsprefixes['w']+'}gridCol':
			lyx_col = lyxTableCol()
			lyx_col.width = table_col.attrib['{'+docx.nsprefixes['w']+'}w']
			lyx_table.columns.append(lyx_col)
	
	# Calculate the total number of columns and add to table properties
	lyx_table.numColumns = len(lyx_table.columns)
	return lyx_table
	

def processTableRow(table_row):
	'''Parse table rows. Create table cells, populate with cell properties.'''
	
	# Create table row structure
	lyx_row = lyxTableRow()
	
	# Iterate through the structures for each row, parse properties, locate cells
	for row_sub in table_row.iter():
		
		# Iterate over table cell, find data and properties
		if row_sub.tag == '{'+docx.nsprefixes['w']+'}tc':
			
			# Create table cell, find properties and data
			lyx_cell = lyxTableCell()
			for cell_prop in row_sub.iter():
				
				# Find cell properties, Add to table cell
				if cell_prop.tag == '{'+docx.nsprefixes['w']+'}tcPr':
					for cell_aspect in cell_prop.iter():

						# Cell Width
						if cell_aspect.tag == '{'+docx.nsprefixes['w']+'}tcW':
							lyx_cell.width = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}w']
							lyx_cell.units = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}type']

						# Mutli-column Cells
						if cell_aspect.tag == '{'+docx.nsprefixes['w']+'}gridSpan':
							lyx_cell.span_multicol = 'true'
							lyx_cell.multicol = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}val']

						# Multi-row cells
						if cell_aspect.tag == '{'+docx.nsprefixes['w']+'}vMerge':
							lyx_cell.span_multirow = 'true'
							try:
								lyx_cell.multirow_start = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}val']
							except:
								pass
				
				# Add cell data, translated to a valid LyX structure
				if cell_prop.tag == '{'+docx.nsprefixes['w']+'}p':
					lyx_cell.data = processParagraph(cell_prop, True)
					
					# Parse paragraph properties for text alignment
					for cell_aspect in cell_prop.iter():
						if cell_aspect.tag == '{'+docx.nsprefixes['w']+'}jc':
							lyx_cell.textalign = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}val']

				# Determine cell borders
				if cell_prop.tag == '{'+docx.nsprefixes['w']+'}tcBorders':
					# Iterate through the cell borders and set lyxCell.borders.data
					for border in cell_prop:

						# Top Border
						if border.tag == '{'+docx.nsprefixes['w']+'}top':
							lyx_cell.borders.data['top'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

						# Bottom Border
						if border.tag == '{'+docx.nsprefixes['w']+'}bottom':
							lyx_cell.borders.data['bottom'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

						# Left Border
						if border.tag == '{'+docx.nsprefixes['w']+'}left':
							lyx_cell.borders.data['left'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

						# Right Border
						if border.tag == '{'+docx.nsprefixes['w']+'}right':
							lyx_cell.borders.data['right'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

			# Add the cell to the row
			lyx_row.tableCells.append(lyx_cell)
						
	return lyx_row


def processStyleAttributes(charstyle, doc_styles, doc_fonts):
	'''Look at the document style information and font information to find relevant'''
	
	# Create Property Lists
	shape = []
	series = []
	misc = []
	
	# Locate the info for the charstyle
	for style in doc_styles.iterchildren():
		if style.tag == '{'+docx.nsprefixes['w']+'}style':
			# Parse the general properties for the style
			if style.attrib['{'+docx.nsprefixes['w']+'}styleId'] == charstyle.styleName:
				
				# Iterate through the style properties
				for style_prop in style.iterchildren():

					# Retrieve the Parent Style Information
					if style_prop.tag == '{'+docx.nsprefixes['w']+'}basedOn':
						charstyle.fontData['ParentStyle'] = style_prop.attrib['{'+docx.nsprefixes['w']+'}val']

					# Look at the run properties (rPr) to determine text shape and series
					if style_prop.tag == '{'+docx.nsprefixes['w']+'}rPr':

						# Iterate through property elements
						for rPr in style_prop.iterchildren():

							# Bold Weight
							if rPr.tag == '{'+docx.nsprefixes['w']+'}b':
								if checkStyleAttrib(rPr.attrib): series.append('Bold')
							
							# Italic
							if rPr.tag == '{'+docx.nsprefixes['w']+'}i':
								if checkStyleAttrib(rPr.attrib): shape.append('Italic')

							# Uppercase
							if rPr.tag == '{'+docx.nsprefixes['w']+'}caps':
								if checkStyleAttrib(rPr.attrib): shape.append('Up')

							# Small Caps
							if rPr.tag == '{'+docx.nsprefixes['w']+'}smallCaps':
								if checkStyleAttrib(rPr.attrib): shape.append('SmallCaps')

							# Font Properties
							if rPr.tag == '{'+docx.nsprefixes['w']+'}rFonts':
								if '{'+docx.nsprefixes['w']+'}cs' in rPr.attrib.keys():
									charstyle.fontData['fontname'] = rPr.attrib['{'+docx.nsprefixes['w']+'}cs']
								elif '{'+docx.nsprefixes['w']+'}ascii' in rPr.attrib.keys():
									charstyle.fontData['fontname'] = rPr.attrib['{'+docx.nsprefixes['w']+'}ascii']
								elif '{'+docx.nsprefixes['w']+'}hAnsi' in rPr.attrib.keys():
									charstyle.fontData['fontname'] = rPr.attrib['{'+docx.nsprefixes['w']+'}hAnsi']
								else:
									charstyle.fontData['fontname'] = ''

	# Look for font characteristsics: family, size
	if charstyle.fontData['fontname'] is not '':

		# Iterate through document fonts to find match
		for font in doc_fonts:
			if charstyle.fontData['fontname'] == font.attrib['{'+docx.nsprefixes['w']+'}name']:
				
				# Parse properties looking for the font-family
				for fontprop in font.iterchildren():
					
					# Translate the family classification from Word to LyX
					if fontprop.tag == '{'+docx.nsprefixes['w']+'}family':
						charstyle.fontData['lyx_family'] = \
							FONT_FAMILIES[fontprop.attrib['{'+docx.nsprefixes['w']+'}val']]

	charstyle.fontData['lyx_series'] = ', '.join(series)
	charstyle.fontData['lyx_shape'] = ', '.join(shape)
	return charstyle

def checkStyleAttrib(attrib):
	'''Check style attributes to determine if a value is disable (val == 0). If disabled, 
	return False. Otherwise, return True.'''
	# Check to see if there are attributes
	if len(attrib) > 0:
		# If the attrib is disabled (val == 0), return false
		if attrib['{'+docx.nsprefixes['w']+'}val'] == '0': return False
	else: return True


#------------ Test Logic ------------## 

lyxoutput = codecs.open(outputfile, 'w', 'utf-8')
doc_body = processDocument(inputfile, outputfile)
lyxoutput.write(doc_body)
lyxoutput.close()

print 'Conversion successful, output written to ' + outputfile
