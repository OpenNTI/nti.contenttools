#!/usr/bin/python

# word2lyx is a document parsing script used to 
# convert Microsoft Word documents to LyX documents.
# (C) Robert Oakes, 2012. Released under the terms
# of the GNU Lesser General Public License (LGPL).
# Preferred Dependencies: lxml, elyxer

# Import docx library and supporting tools
import sys, os, tempfile, shutil, codecs

# Import docx parsing classes
from docx import read as docxread
from docx import properties as docx

# Import lyx parsing classes
from lyx.tables import *
from lyx.charstyles import *
import lyx.templates as templates

##------------ Setup Input File, Output File, and Template ------------## 
SCRIPT_DIR = os.path.abspath(os.path.dirname(__file__))
template = os.path.join(SCRIPT_DIR, 'lyx/templates/article.w2l')

# Get input and output file
try:
	inputfile = sys.argv[1]
	# Check to see if the file exists
	if not os.path.exists(inputfile):
		print 'Error: Unable to locate input file'
		exit()
	# Check to see if the file is a docx file
	basename = os.path.basename(inputfile)
	fileext = basename.split('.')[len(basename.split('.'))-1]
	if fileext != 'docx':
		print 'Error: Invalid input file. word2lyx only supports docx files.'
		exit()
	outputfile = sys.argv[2]

# If there is an error encoutered, return general error message
except:
	print 'lyx2word encountered an error with the file' 
	exit()

for index, option in enumerate(sys.argv):
	# Set the template that word2lyx should use for the conversion
	if option == '-t':
		template = 'lyx/templates/' + sys.argv[index+1] + '.w2l'
		template = os.path.join(SCRIPT_DIR, template)
		if os.path.exists(template) != True:
			print 'Please specify a valid template'
			exit()


##------------ Retrieve Template ------------##
doc_options = templates.DocumentOptions(template)


##------------ Global Structures ------------## 

CHARSTYLE_LIST = []
PROC_REL = []
DOC_REL = docxread.openDocxRelationships(inputfile)

# Try to load footnotes
try:
	DOC_FOOTNOTES = docxread.openDocxFootnotes(inputfile)
except:
	DOC_FOOTNOTES = None

# Try to load endnotes

try:
	DOC_ENDNOTES = docxread.openDocxEndnotes(inputfile)
except:
	DOC_ENDNOTES = None

PARAGRAPH_STYLES = doc_options.returnStyles('ParagraphStyles')
TABLE_STYLES = doc_options.returnStyles('TableStyles')
CHARACTER_STYLES = doc_options.returnStyles('CharacterStyles')
IGNORE_STYLES = doc_options.returnSetting('IgnoreStyles')

IMAGE_FOLDER = doc_options.returnSetting('ImageDir')


##------------ Class Methods ------------## 

def processDocument(inputfile, outputfile, doc_options):
	'''Method that parses the input document and translates it to a valid LyX structure. 
		Returns a text string with XML structures translated to equivalent LyX insets. 
		It uses the layouts in article and book as a base. If there are character styles 
		not defined in article or book, it will create placeholder styles using flex:inset.
		These may be modified at a later point by the user.'''

	# Open document from input string
	print "Beginning Conversion of " + inputfile
	document = docxread.opendocx(inputfile)

	# Create output strings
	doc_header = u''
	doc_layout = u''
	doc_body = u''
	
	# Iterate over the structure of the document, process document body
	for element in document.iterchildren():
		
		# Process Elements in Document Body
		if element.tag == '{'+docx.nsprefixes['w']+'}body':
			doc_body = doc_body + u'\\begin_body\n'
			for body_element in element.iterchildren():
				out_text = u''
				
				# P (paragraph) Elements
				if body_element.tag == '{'+docx.nsprefixes['w']+'}p':
					out_text = processParagraph(body_element)
			
				if body_element.tag == '{'+docx.nsprefixes['w']+'}tbl':
					out_text = processTable(body_element)
		
				# Write to the Document Body
				if len(out_text) > 0:
					doc_body = doc_body + out_text
					
			doc_body = doc_body + u'\\end_body\n'

	# Copy Document Images to Subfolder
	if len(PROC_REL) > 0:
		img_exportfolder = os.path.join(os.path.abspath(os.path.dirname(outputfile)),
			IMAGE_FOLDER)
		docxread.getDocumentImages(inputfile, PROC_REL, img_exportfolder)
	
	# Retrieve character style information for document
	doc_styles = docxread.openDocxStyles(inputfile)
	# Retrieve font information for document
	doc_fonts = docxread.openFontTable(inputfile)

	# Add character styles in the list to the document local layout
	for charstyle in CHARSTYLE_LIST:
		charstyle = processStyleAttributes(charstyle, doc_styles, doc_fonts)
		charstyle_def = writeLyxCharstyle(charstyle)
		doc_layout = doc_layout + '\n' + charstyle_def

	# Add opening and closing tags to Local Layout
	if len(doc_layout) > 0:
		doc_layout = '\\begin_local_layout' + doc_layout + '\n\\end_local_layout'

	# Create the document header
	for option_section in doc_options.optionSections():
		# Retrieve the options for a particular section
		options = doc_options.getDocOptions(option_section)

		# Document Default Options
		if option_section == 'DocumentDefaults':
			for option in options:
				# Look for first line, write comment line
				if option == 'firstline':
					doc_header = doc_header + options[option] + '\n'
				elif option == '\\lyxformat':
					doc_header = doc_header + option + ' ' + options[option] \
						+ '\n\\begin_document\n\\begin_header\n'
				elif option == 'modules':
					doc_header = doc_header + '\\begin_modules\n' \
						+ options[option] + '\n\\end_modules\n'
				else:
					doc_header = doc_header + option + ' ' + options[option] + '\n'

			if len(doc_layout) > 0:
				doc_header = doc_header + doc_layout + '\n'

		# LaTeX Package Options
		elif option_section == 'PackageOptions':
			for option in options:
				doc_header = doc_header + '\\use_package ' + option \
					+ ' ' + options[option] + '\n'

		elif option_section == 'IndexOptions':
			doc_header = doc_header + '\\index Index\n'
			for option in options:
				doc_header = doc_header + option + ' ' + options[option] + '\n'
			doc_header = doc_header + '\\end_index\n'

		# Citation Options, Page Options, Index Options, TOCOptions
		else:
			for option in options:
				doc_header = doc_header + option + ' ' + options[option] + '\n'
	doc_header = doc_header + '\\end_header\n\n'

	# Combine the document header and body
	lyx_doc = doc_header + doc_body + '\n\\end_document\n'

	return lyx_doc

def processParagraph(paragraph, lyxTable = False):
	'''Processes the text of a given paragraph into insets and text.'''
	
	partext = u''
	parstyle = u''
	lyx_paragraph = u''
		
	# Scan the elements in the paragraph and extract information
	for element in paragraph.iterchildren():

		# Process text runs
		# Process Text Runs
		if element.tag == '{'+docx.nsprefixes['w']+'}r':
			partext = processTextRun(element, partext)
			image_text = processRelationship(element)
			if len(image_text) > 0:
				partext = partext + image_text
		
		# Paragraph Properties
		if element.tag == '{'+docx.nsprefixes['w']+'}pPr':
			for sub_element in element.iterchildren():
				# Look for Paragraph Styles
				if sub_element.tag == '{'+docx.nsprefixes['w']+'}pStyle':
					parstyle = sub_element.attrib['{'+docx.nsprefixes['w']+'}val']
	
	# Write Inset
	lyxstylename = 'Standard'
	if len(partext) > 0:
		
		# Check for standard paragraph styles
		if (len(parstyle) > 0) & (lyxTable == False):
			if parstyle in PARAGRAPH_STYLES.keys():
				lyxstylename = PARAGRAPH_STYLES[parstyle]
		elif (len(parstyle) > 0) & (lyxTable == True):
			if parstyle in TABLE_STYLES.keys():
				lyxstylename = TABLE_STYLES[parstyle]
		
		lyx_paragraph = '\\begin_layout ' + lyxstylename + u'\n' + partext \
			+ u'\n\\end_layout\n'
		
		if lyxTable == True:
			lyx_paragraph = '\\begin_inset Text\n' + lyx_paragraph \
			+ '\\end_inset\n'
	
	return lyx_paragraph


def processTextRun(textrun, partext):
	'''Process a paragraph textrun, parse for character styles'''

	rstyle = False
	rstyle_text = u''
	rstyle_name = u''

	for element in textrun.iterchildren():

		# Look for character styles inside of the run
		if element.tag == '{'+docx.nsprefixes['w']+'}rPr':
			rstyle = True
			for sub_element in element.iterchildren():
				if sub_element.tag == '{'+docx.nsprefixes['w']+'}rStyle':
					rstyle_name = sub_element.attrib['{'+docx.nsprefixes['w']+'}val']

					# Compare Style Name to Ignore List
					if rstyle_name in IGNORE_STYLES:
						rstyle = False

		# Find paragraph text
		if (element.tag == '{'+docx.nsprefixes['w']+'}t'): 
			# If not character style, append to the end of the paragraph
			if element.text:
				if rstyle == False:
					partext = partext + element.text
				elif rstyle == True:
					rstyle_text = rstyle_text + element.text

		# Look for footnotes
		if (element.tag == '{'+docx.nsprefixes['w']+'}footnoteReference'):
			if rstyle == False:
				partext = processFootnote(element, partext)
			elif rstyle == True:
				rstyle_text = processFootnote(element, rstyle_text)

		# Look for other runs embedded in this run, process recursively
		if (element.tag == '{'+docx.nsprefixes['w']+'}r'):
			if rstyle == False:
				partext = processTextRun(element, partext)
			elif rstyle == True:
				rstyle_text = processTextRun(element, rstyle_text)

	# For character styles, create a flex inset
	if (rstyle == True) & (len(rstyle_text) > 0):
		# Check defined character styles
		if rstyle_name in CHARACTER_STYLES.keys():
			lyxcstyle = CHARACTER_STYLES[rstyle_name]

		# Return
		elif rstyle_name == '':
			return partext + rstyle_text

		# If not defined, generate a skeleton style and add to list
		else:
			lyxcstyle = rstyle_name
			cstyle = lyxCharacterStyle(rstyle_name)

			# Check to see if the style has already been added
			duplicates = [style for style in CHARSTYLE_LIST if style.styleName == rstyle_name]
			if len(duplicates) == 0:
				CHARSTYLE_LIST.append(cstyle)

		partext = partext + '\n\\begin_inset Flex ' + lyxcstyle + '\nstatus collapsed\n' + \
			'\\begin_layout Plain Layout\n'+ rstyle_text + '\n\\end_layout\n\\end_inset\n'

	return partext


def processFootnote(footnote_ref, partext):
	'''Locate footnote and write a footnote inset.'''

	# Retrieve the footnote Id Number
	foot_id = footnote_ref.attrib['{'+docx.nsprefixes['w']+'}id']
	foot_text = u''

	# Retrieve the footnote text
	for footnote in DOC_FOOTNOTES.iterchildren():
		if footnote.attrib['{'+docx.nsprefixes['w']+'}id'] == foot_id:
			for foot_sub in footnote.iterchildren():
				# Process paragraphs found in the note
				if foot_sub.tag == '{'+docx.nsprefixes['w']+'}p':
					foot_text = processParagraph(foot_sub, foot_text)

	partext = partext + '\n\\begin_inset Foot\nstatus collapsed\n' \
		+ foot_text + '\\end_inset\n'

	return partext


def processRelationship(relationship):
	'''Process relationships in the document. '''
	rel_text = u''
	
	# Iterate through the relationship properties, process
	for element in relationship.iter():
		if element.tag == '{'+docx.nsprefixes['w']+'}drawing':
			rel_text = processImage(element)
	return rel_text


def processImage(image):
	'''Process images in the document. Specify that image should be copied from the 
		zip archive to a destination directory. Add a LyX image inset to the document. 
		The method ignores Microsoft specific charts and other files that have been 
		embedded.'''

	image_text = u''
	rId = u''
	rel_type = u''
	rel_name = u''
	rel_target = u''
	rel_props = u'width 100text%'

	# Iterate through the image properties, process
	for element in image.iter():
		
		# Search for Image Properties, contained in blipFill
		if element.tag == '{'+docx.nsprefixes['pic']+'}blipFill':
			# Retrieve the Image rId and filename
			for imageprop in element.iter():
				if imageprop.tag == '{'+docx.nsprefixes['a']+'}blip':
					rId = imageprop.attrib['{'+docx.nsprefixes['r']+'}embed']
					rel_type, rel_target = relationshipProperties(rId)
					PROC_REL.append(os.path.basename(rel_target))
					break
	
	# Write inset to file, ignore charts and other types of files
	if rel_type == docx.nsprefixes['wordimage']:
		image_text = '\\begin_inset Graphics\n' + '\tfilename ' + IMAGE_FOLDER + \
			'/' + os.path.basename(rel_target) + '\n' + u'\t' + rel_props + \
			u'\n' + u'\\end_inset\n'
	
	return image_text


def relationshipProperties(rId):
	'''Parse the document relationships to retrieve the type and target for a 
		specified document element.'''

	# Search through DOC_REL to retrieve relationship properties
	for relationship in DOC_REL.iter():
		try:
			if relationship.attrib['Id'] == rId:
				return relationship.attrib['Type'], relationship.attrib['Target']
		except:
			pass


def processTable(table_element):
	'''Parse a table element and translate to LyX.'''
	# Create a logical representation of the lyxTable
	lyx_table = lyxTable()
	
	# Parse element and determine table properties
	for element in table_element.iterchildren():

		# Parse table properties
		if element.tag == '{'+docx.nsprefixes['w']+'}tblPr':
			lyx_table = tableProperties(element, lyx_table)
				
		# Create column properties, add to lyxTable
		if element.tag == '{'+docx.nsprefixes['w']+'}tblGrid':
			lyx_table = tableGridProperties(element, lyx_table)
			
		# Create table rows, cells. Populate with data and properties.
		# Add populated rows to the lyxTable
		if element.tag == '{'+docx.nsprefixes['w']+'}tr':
			lyx_row = processTableRow(element)
			lyx_table.rows.append(lyx_row)
	
	# Add additional data to table structure
	lyx_table.numRows = len(lyx_table.rows)
	out_text = writeLyxTable(lyx_table)
	out_text = '\\begin_layout Standard\n' + out_text + '\\end_layout\n'
	
	return out_text


def tableProperties(property_list, lyx_table):
	'''Parse the general table properties, such as determining the location of table 
		borders, whether there is a grid in use, and the width of the table.'''

	# Scan the tblProperty elements
	for prop in property_list.iter():

		# Look for general table style settings
		if prop.tag == '{'+docx.nsprefixes['w']+'}tblStyle':
			if '{'+docx.nsprefixes['w']+'}val' in prop.attrib.keys():
				# Apply a grid to the table, activate all borders
				if prop.attrib['{'+docx.nsprefixes['w']+'}val'] == 'TableGrid':
					for border in lyx_table.borders.data.keys():
						lyx_table.borders.data[border] = 'true'

		# Look for table border settings
		if prop.tag == '{'+docx.nsprefixes['w']+'}tblBorders':

			# Apply settings to the border settings for the lyxTable object
			for border in prop.iterchildren():

				# Top Border
				if border.tag == '{'+docx.nsprefixes['w']+'}top':
					lyx_table.borders.data['top'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

				# Bottom Border
				if border.tag == '{'+docx.nsprefixes['w']+'}bottom':
					lyx_table.borders.data['bottom'] = border.attrib['{'+docx.nsprefixes['w']+'}val']	
				# Left Border
				if border.tag == '{'+docx.nsprefixes['w']+'}left':
					lyx_table.borders.data['left'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

				# Right Border
				if border.tag == '{'+docx.nsprefixes['w']+'}right':
					lyx_table.borders.data['right'] = border.attrib['{'+docx.nsprefixes['w']+'}val']
				
				# Inside Horizontal Borders
				if border.tag == '{'+docx.nsprefixes['w']+'}insideH':
					lyx_table.borders.data['insideH'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

				# Inside Veritcal Border
				if border.tag == '{'+docx.nsprefixes['w']+'}insideV':
					lyx_table.borders.data['insideV'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

		# Table Width and Units
		if prop.tag == '{'+docx.nsprefixes['w']+'}tblW':
			lyx_table.width = prop.attrib['{'+docx.nsprefixes['w']+'}w']
			lyx_table.units = prop.attrib['{'+docx.nsprefixes['w']+'}type']

	return lyx_table
	

def tableGridProperties(table_grid, lyx_table):
	'''Parse grid column properties. Looks at the table column properties to set width and text styles. Takes both the element to be parsed and a lyxTable object as input. Returns the lyxTable with column objects appended.'''
	
	# Scan the tblGrid Element for gridCol elements, retrieve relative width for each
	for table_col in table_grid.iter():
		if table_col.tag == '{'+docx.nsprefixes['w']+'}gridCol':
			lyx_col = lyxTableCol()
			lyx_col.width = table_col.attrib['{'+docx.nsprefixes['w']+'}w']
			lyx_table.columns.append(lyx_col)
	
	# Calculate the total number of columns and add to table properties
	lyx_table.numColumns = len(lyx_table.columns)
	return lyx_table
	

def processTableRow(table_row):
	'''Parse table rows. Create table cells, populate with cell properties.'''
	
	# Create table row structure
	lyx_row = lyxTableRow()
	
	# Iterate through the structures for each row, parse properties, locate cells
	for row_sub in table_row.iter():
		
		# Iterate over table cell, find data and properties
		if row_sub.tag == '{'+docx.nsprefixes['w']+'}tc':
			
			# Create table cell, find properties and data
			lyx_cell = lyxTableCell()
			for cell_prop in row_sub.iter():
				
				# Find cell properties, Add to table cell
				if cell_prop.tag == '{'+docx.nsprefixes['w']+'}tcPr':
					for cell_aspect in cell_prop.iter():

						# Cell Width
						if cell_aspect.tag == '{'+docx.nsprefixes['w']+'}tcW':
							lyx_cell.width = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}w']
							lyx_cell.units = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}type']

						# Mutli-column Cells
						if cell_aspect.tag == '{'+docx.nsprefixes['w']+'}gridSpan':
							lyx_cell.span_multicol = 'true'
							lyx_cell.multicol = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}val']

						# Multi-row cells
						if cell_aspect.tag == '{'+docx.nsprefixes['w']+'}vMerge':
							lyx_cell.span_multirow = 'true'
							try:
								lyx_cell.multirow_start = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}val']
							except:
								pass
				
				# Add cell data, translated to a valid LyX structure
				if cell_prop.tag == '{'+docx.nsprefixes['w']+'}p':
					lyx_cell.data = processParagraph(cell_prop, True)
					
					# Parse paragraph properties for text alignment
					for cell_aspect in cell_prop.iter():
						if cell_aspect.tag == '{'+docx.nsprefixes['w']+'}jc':
							lyx_cell.textalign = cell_aspect.attrib['{'+docx.nsprefixes['w']+'}val']

				# Determine cell borders
				if cell_prop.tag == '{'+docx.nsprefixes['w']+'}tcBorders':
					# Iterate through the cell borders and set lyxCell.borders.data
					for border in cell_prop:

						# Top Border
						if border.tag == '{'+docx.nsprefixes['w']+'}top':
							lyx_cell.borders.data['top'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

						# Bottom Border
						if border.tag == '{'+docx.nsprefixes['w']+'}bottom':
							lyx_cell.borders.data['bottom'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

						# Left Border
						if border.tag == '{'+docx.nsprefixes['w']+'}left':
							lyx_cell.borders.data['left'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

						# Right Border
						if border.tag == '{'+docx.nsprefixes['w']+'}right':
							lyx_cell.borders.data['right'] = border.attrib['{'+docx.nsprefixes['w']+'}val']

			# Add the cell to the row
			lyx_row.tableCells.append(lyx_cell)
						
	return lyx_row


def processStyleAttributes(charstyle, doc_styles, doc_fonts):
	'''Look at the document style information and font information to find relevant'''
	
	# Create Property Lists
	shape = []
	series = []
	misc = []
	
	# Locate the info for the charstyle
	for style in doc_styles.iterchildren():
		if style.tag == '{'+docx.nsprefixes['w']+'}style':
			# Parse the general properties for the style
			if style.attrib['{'+docx.nsprefixes['w']+'}styleId'] == charstyle.styleName:
				
				# Iterate through the style properties
				for style_prop in style.iterchildren():

					# Retrieve the Parent Style Information
					if style_prop.tag == '{'+docx.nsprefixes['w']+'}basedOn':
						charstyle.fontData['ParentStyle'] = style_prop.attrib['{'+docx.nsprefixes['w']+'}val']

					# Look at the run properties (rPr) to determine text shape and series
					if style_prop.tag == '{'+docx.nsprefixes['w']+'}rPr':

						# Iterate through property elements
						for rPr in style_prop.iterchildren():

							# Bold Weight
							if rPr.tag == '{'+docx.nsprefixes['w']+'}b':
								if checkStyleAttrib(rPr.attrib): series.append('Bold')
							
							# Italic
							if rPr.tag == '{'+docx.nsprefixes['w']+'}i':
								if checkStyleAttrib(rPr.attrib): shape.append('Italic')

							# Uppercase
							if rPr.tag == '{'+docx.nsprefixes['w']+'}caps':
								if checkStyleAttrib(rPr.attrib): shape.append('Up')

							# Small Caps
							if rPr.tag == '{'+docx.nsprefixes['w']+'}smallCaps':
								if checkStyleAttrib(rPr.attrib): shape.append('SmallCaps')

							# Font Properties
							if rPr.tag == '{'+docx.nsprefixes['w']+'}rFonts':
								if '{'+docx.nsprefixes['w']+'}cs' in rPr.attrib.keys():
									charstyle.fontData['fontname'] = rPr.attrib['{'+docx.nsprefixes['w']+'}cs']
								elif '{'+docx.nsprefixes['w']+'}ascii' in rPr.attrib.keys():
									charstyle.fontData['fontname'] = rPr.attrib['{'+docx.nsprefixes['w']+'}ascii']
								elif '{'+docx.nsprefixes['w']+'}hAnsi' in rPr.attrib.keys():
									charstyle.fontData['fontname'] = rPr.attrib['{'+docx.nsprefixes['w']+'}hAnsi']
								else:
									charstyle.fontData['fontname'] = ''

	# Look for font characteristsics: family, size
	if charstyle.fontData['fontname'] is not '':

		# Iterate through document fonts to find match
		for font in doc_fonts:
			if charstyle.fontData['fontname'] == font.attrib['{'+docx.nsprefixes['w']+'}name']:
				
				# Parse properties looking for the font-family
				for fontprop in font.iterchildren():
					
					# Translate the family classification from Word to LyX
					if fontprop.tag == '{'+docx.nsprefixes['w']+'}family':
						charstyle.fontData['lyx_family'] = \
							FONT_FAMILIES[fontprop.attrib['{'+docx.nsprefixes['w']+'}val']]

	charstyle.fontData['lyx_series'] = ', '.join(series)
	charstyle.fontData['lyx_shape'] = ', '.join(shape)
	return charstyle

def checkStyleAttrib(attrib):
	'''Check style attributes to determine if a value is disable (val == 0). If disabled, 
	return False. Otherwise, return True.'''
	# Check to see if there are attributes
	if len(attrib) > 0:
		# If the attrib is disabled (val == 0), return false
		if attrib['{'+docx.nsprefixes['w']+'}val'] == '0': return False
	else: return True


#------------ Test Logic ------------## 

lyxoutput = codecs.open(outputfile, 'w', 'utf-8')
doc_body = processDocument(inputfile, outputfile, doc_options)
lyxoutput.write(doc_body)
lyxoutput.close()

print 'Conversion successful, output written to ' + outputfile